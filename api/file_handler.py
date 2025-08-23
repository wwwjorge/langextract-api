"""File handling module for LangExtract API."""

import os
import uuid
import mimetypes
from typing import Optional, List
from pathlib import Path

from fastapi import UploadFile, HTTPException, status
import aiofiles
import structlog

from api.config import get_settings

logger = structlog.get_logger()


class FileHandler:
    """Handle file uploads and processing."""
    
    def __init__(self):
        self.settings = get_settings()
        self._ensure_directories()
    
    def _ensure_directories(self):
        """Ensure upload and output directories exist."""
        os.makedirs(self.settings.uploads_dir, exist_ok=True)
        os.makedirs(self.settings.outputs_dir, exist_ok=True)
    
    async def process_upload(self, file: UploadFile) -> str:
        """Process uploaded file and return text content."""
        # Validate file
        await self._validate_file(file)
        
        # Save file temporarily
        file_path = await self._save_upload(file)
        
        try:
            # Extract text content
            content = await self._extract_text_content(file_path, file.content_type)
            
            logger.info(
                "File processed successfully",
                filename=file.filename,
                content_length=len(content),
                content_type=file.content_type
            )
            
            return content
            
        finally:
            # Clean up temporary file
            try:
                os.remove(file_path)
            except Exception as e:
                logger.warning("Failed to remove temporary file", file_path=file_path, error=str(e))
    
    async def _validate_file(self, file: UploadFile):
        """Validate uploaded file."""
        if not file.filename:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="No filename provided"
            )
        
        # Check file extension
        file_ext = Path(file.filename).suffix.lower()
        allowed_extensions = self.settings.allowed_file_extensions_list
        
        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File type {file_ext} not allowed. Allowed types: {', '.join(allowed_extensions)}"
            )
        
        # Check file size
        file.file.seek(0, 2)  # Seek to end
        file_size = file.file.tell()
        file.file.seek(0)  # Reset to beginning
        
        if file_size > self.settings.max_file_size_bytes:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail=f"File too large. Maximum size: {self.settings.max_file_size_mb}MB"
            )
        
        logger.info(
            "File validation passed",
            filename=file.filename,
            size=file_size,
            content_type=file.content_type
        )
    
    async def _save_upload(self, file: UploadFile) -> str:
        """Save uploaded file temporarily."""
        file_id = str(uuid.uuid4())
        file_ext = Path(file.filename).suffix.lower()
        temp_filename = f"{file_id}{file_ext}"
        file_path = os.path.join(self.settings.uploads_dir, temp_filename)
        
        try:
            async with aiofiles.open(file_path, 'wb') as f:
                content = await file.read()
                await f.write(content)
            
            logger.info(
                "File saved temporarily",
                original_filename=file.filename,
                temp_path=file_path,
                size=len(content)
            )
            
            return file_path
            
        except Exception as e:
            logger.error(
                "Failed to save uploaded file",
                filename=file.filename,
                error=str(e)
            )
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Failed to save uploaded file"
            )
    
    async def _extract_text_content(self, file_path: str, content_type: Optional[str]) -> str:
        """Extract text content from file based on its type."""
        file_ext = Path(file_path).suffix.lower()
        
        try:
            if file_ext == '.txt' or file_ext == '.md':
                return await self._read_text_file(file_path)
            elif file_ext == '.json':
                return await self._read_json_file(file_path)
            elif file_ext == '.pdf':
                return await self._read_pdf_file(file_path)
            elif file_ext == '.docx':
                return await self._read_docx_file(file_path)
            else:
                # Try to read as text file
                return await self._read_text_file(file_path)
                
        except Exception as e:
            logger.error(
                "Failed to extract text content",
                file_path=file_path,
                file_ext=file_ext,
                error=str(e)
            )
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Failed to extract text from {file_ext} file: {str(e)}"
            )
    
    async def _read_text_file(self, file_path: str) -> str:
        """Read plain text file."""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                async with aiofiles.open(file_path, 'r', encoding=encoding) as f:
                    content = await f.read()
                    logger.info("Text file read successfully", encoding=encoding)
                    return content
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Could not decode text file with any supported encoding")
    
    async def _read_json_file(self, file_path: str) -> str:
        """Read JSON file and convert to text."""
        import json
        
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            content = await f.read()
            
        try:
            # Parse JSON to validate
            data = json.loads(content)
            
            # Convert to formatted text
            if isinstance(data, dict):
                text_parts = []
                for key, value in data.items():
                    text_parts.append(f"{key}: {json.dumps(value, ensure_ascii=False)}")
                return "\n".join(text_parts)
            elif isinstance(data, list):
                text_parts = []
                for i, item in enumerate(data):
                    text_parts.append(f"Item {i+1}: {json.dumps(item, ensure_ascii=False)}")
                return "\n".join(text_parts)
            else:
                return str(data)
                
        except json.JSONDecodeError:
            # If not valid JSON, return as text
            return content
    
    async def _read_pdf_file(self, file_path: str) -> str:
        """Read PDF file and extract text."""
        try:
            import PyPDF2
            
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_content.append(f"--- Page {page_num + 1} ---\n{text}")
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}", error=str(e))
            
            if not text_content:
                raise ValueError("No text content found in PDF")
            
            return "\n\n".join(text_content)
            
        except ImportError:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="PDF processing not available. PyPDF2 library not installed."
            )
    
    async def _read_docx_file(self, file_path: str) -> str:
        """Read DOCX file and extract text."""
        try:
            import docx
            
            doc = docx.Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            if not text_content:
                raise ValueError("No text content found in DOCX")
            
            return "\n\n".join(text_content)
            
        except ImportError:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="DOCX processing not available. python-docx library not installed."
            )
    
    def get_file_info(self, file_path: str) -> dict:
        """Get information about a file."""
        if not os.path.exists(file_path):
            return None
        
        stat = os.stat(file_path)
        mime_type, _ = mimetypes.guess_type(file_path)
        
        return {
            "path": file_path,
            "size": stat.st_size,
            "created": stat.st_ctime,
            "modified": stat.st_mtime,
            "mime_type": mime_type,
            "extension": Path(file_path).suffix.lower()
        }
    
    def list_uploaded_files(self) -> List[dict]:
        """List all uploaded files."""
        files = []
        
        try:
            for filename in os.listdir(self.settings.uploads_dir):
                file_path = os.path.join(self.settings.uploads_dir, filename)
                if os.path.isfile(file_path):
                    file_info = self.get_file_info(file_path)
                    if file_info:
                        file_info["filename"] = filename
                        files.append(file_info)
        
        except Exception as e:
            logger.error("Failed to list uploaded files", error=str(e))
        
        return files
    
    def cleanup_old_files(self, days: int = 7):
        """Clean up old uploaded files."""
        from datetime import datetime, timedelta
        
        cutoff_time = datetime.now() - timedelta(days=days)
        removed_count = 0
        
        try:
            for filename in os.listdir(self.settings.uploads_dir):
                file_path = os.path.join(self.settings.uploads_dir, filename)
                if os.path.isfile(file_path):
                    file_time = datetime.fromtimestamp(os.path.getctime(file_path))
                    if file_time < cutoff_time:
                        os.remove(file_path)
                        removed_count += 1
                        logger.info("Removed old uploaded file", file_path=file_path)
        
        except Exception as e:
            logger.error("Failed to cleanup old uploaded files", error=str(e))
        
        logger.info("Upload cleanup completed", removed_files=removed_count)
        return removed_count
    
    async def save_extraction_result(self, extraction_id: str, content: str, format: str = "txt") -> str:
        """Save extraction result to file."""
        filename = f"{extraction_id}_result.{format}"
        file_path = os.path.join(self.settings.outputs_dir, filename)
        
        try:
            async with aiofiles.open(file_path, 'w', encoding='utf-8') as f:
                await f.write(content)
            
            logger.info(
                "Extraction result saved",
                extraction_id=extraction_id,
                file_path=file_path,
                format=format
            )
            
            return file_path
            
        except Exception as e:
            logger.error(
                "Failed to save extraction result",
                extraction_id=extraction_id,
                error=str(e)
            )
            raise
    
    def get_supported_formats(self) -> dict:
        """Get information about supported file formats."""
        return {
            "input_formats": {
                ".txt": "Plain text files",
                ".md": "Markdown files",
                ".json": "JSON files",
                ".pdf": "PDF documents (requires PyPDF2)",
                ".docx": "Microsoft Word documents (requires python-docx)"
            },
            "output_formats": {
                "json": "JSON format",
                "jsonl": "JSON Lines format",
                "html": "HTML visualization",
                "txt": "Plain text format"
            },
            "max_file_size_mb": self.settings.max_file_size_mb,
            "allowed_extensions": self.settings.allowed_file_extensions_list
        }